import logging
import os
import shutil
import time
from typing import BinaryIO, Dict, Any, Optional

from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import VlmPipelineOptions
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.pipeline.vlm_pipeline import VlmPipeline
from docling.datamodel import vlm_model_specs
from pypdf import PdfReader
from docx import Document as DocxDocument

from .settings import settings

logger = logging.getLogger(__name__)


class FileProcessingError(Exception):
    pass


class UnsupportedFileTypeError(FileProcessingError):
    pass


class TextExtractionError(FileProcessingError):
    pass


class DoclingVLMConverter:
    """Docling Converter mit VLM oder Standard Pipeline"""
    _instance: Optional["DoclingVLMConverter"] = None
    _converter = None
    _disabled = False

    @classmethod
    def get_instance(cls) -> Optional["DoclingVLMConverter"]:
        if cls._disabled or not settings.use_docling_parser:
            return None
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def __init__(self):
        if DoclingVLMConverter._disabled:
            return

        init_start = time.time()
        logger.info("🔧 [DOCLING INIT] Initializing Docling converter...")

        try:
            if settings.docling_use_vlm:
                logger.info("   → Mode: VLM Pipeline (GraniteDocling)")
                logger.info("   → Warning: Very slow on CPU (~15-20s per page)")

                pipeline_options = VlmPipelineOptions(
                    vlm_options=getattr(
                        vlm_model_specs,
                        settings.get_vlm_model_spec()
                    ),
                    generate_page_images=settings.docling_generate_images,
                )

                self._converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: PdfFormatOption(
                            pipeline_cls=VlmPipeline,
                            pipeline_options=pipeline_options
                        )
                    }
                )

                init_time = time.time() - init_start
                logger.info(f"✅ [DOCLING INIT] VLM converter ready ({init_time:.1f}s)")
                logger.info(f"   • Model: {settings.docling_vlm_model}")
                logger.info(f"   • Backend: {settings.docling_vlm_backend}")
                logger.info(f"   • Table recognition: VLM-based (~92% accuracy)")

            else:
                logger.info("   → Mode: Standard Pipeline (TableFormer)")
                logger.info("   → Loading RT-DETR layout + TableFormer models...")

                from docling.datamodel.pipeline_options import PdfPipelineOptions

                pipeline_options = PdfPipelineOptions()
                pipeline_options.do_ocr = False
                pipeline_options.do_table_structure = True
                pipeline_options.do_code_enrichment = True
                pipeline_options.generate_page_images = False
                pipeline_options.generate_picture_images = False

                self._converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: PdfFormatOption(
                            pipeline_options=pipeline_options
                        )
                    }
                )

                init_time = time.time() - init_start
                logger.info(f"✅ [DOCLING INIT] Standard converter ready ({init_time:.1f}s)")
                logger.info(f"   • Layout model: RT-DETR v2")
                logger.info(f"   • Table recognition: TableFormer (~85-90% accuracy)")
                logger.info(f"   • OCR: Disabled")

        except Exception as exc:
            init_time = time.time() - init_start
            logger.warning(f"⚠️  [DOCLING INIT] Initialization failed after {init_time:.1f}s: {exc}")
            logger.info("   → Falling back to PyPDF for all conversions")
            DoclingVLMConverter._disabled = True
            self._converter = None

    def convert(self, file_path: str) -> Optional[str]:
        """Konvertiert PDF zu Markdown mit Fortschritts-Logging"""
        if self._converter is None:
            return None

        filename = os.path.basename(file_path)
        convert_start = time.time()

        try:
            try:
                reader = PdfReader(file_path)
                num_pages = len(reader.pages)
                file_size = os.path.getsize(file_path)
                logger.info("=" * 80)
                logger.info(f"📄 [DOCLING] Starting conversion: {filename}")
                logger.info(f"   • File size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
                logger.info(f"   • Total pages: {num_pages}")

                if settings.docling_use_vlm:
                    est_time = num_pages * 15
                    logger.info(f"   • Estimated time: ~{est_time // 60} min {est_time % 60} sec (VLM on CPU)")
                else:
                    est_time = num_pages * 2.5
                    logger.info(f"   • Estimated time: ~{est_time // 60} min {est_time % 60} sec (Standard)")

                logger.info("=" * 80)
            except Exception:
                num_pages = 0
                logger.info(f"📄 [DOCLING] Converting: {filename}")

            # Conversion durchführen
            logger.info("🔄 [DOCLING] Running document converter...")
            logger.info("   → This may take several minutes, please wait...")

            result = self._converter.convert(file_path)
            document = result.document

            if document is None:
                logger.warning(f"⚠️  [DOCLING] No document returned for {filename}")
                return None

            # Export zu Markdown
            logger.info("📝 [DOCLING] Exporting to Markdown...")
            markdown = document.export_to_markdown()

            if markdown:
                convert_time = time.time() - convert_start
                num_tables = len(document.tables)
                num_pictures = len(document.pictures) if hasattr(document, 'pictures') else 0

                logger.info("=" * 80)
                logger.info(f"✅ [DOCLING] Conversion successful!")
                logger.info(f"   • File: {filename}")
                logger.info(f"   • Duration: {convert_time:.1f}s ({convert_time / 60:.1f} min)")

                if num_pages > 0:
                    logger.info(f"   • Speed: {convert_time / num_pages:.1f}s per page")

                logger.info(f"   • Output: {len(markdown):,} characters")
                logger.info(f"   • Tables detected: {num_tables}")
                logger.info(f"   • Pictures: {num_pictures}")
                logger.info("=" * 80)

                return markdown

            logger.warning(f"⚠️  [DOCLING] Empty output for {filename}")
            return None

        except Exception as exc:
            convert_time = time.time() - convert_start
            logger.error("=" * 80)
            logger.error(f"❌ [DOCLING] Conversion failed after {convert_time:.1f}s")
            logger.error(f"   • File: {filename}")
            logger.error(f"   • Error: {type(exc).__name__}: {exc}")
            logger.error("=" * 80)
            return None


class PDFExtractor:
    """PDF Text-Extraktion mit Docling + PyPDF Fallback"""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extrahiert Text aus PDF mit Fortschrittsanzeige
        - Versucht zuerst Docling (falls aktiviert)
        - Fallback zu PyPDF bei Fehler
        """
        filename = os.path.basename(file_path)
        extraction_start = time.time()

        # 1. Versuche Docling
        if settings.use_docling_parser:
            logger.info(f"📄 [EXTRACT] Attempting Docling parser...")
            docling = DoclingVLMConverter.get_instance()

            if docling:
                docling_text = docling.convert(file_path)
                if docling_text:
                    extraction_time = time.time() - extraction_start
                    logger.info(f"✅ [EXTRACT] Docling extraction complete")
                    logger.info(f"   • Duration: {extraction_time:.1f}s")
                    logger.info(f"   • Characters: {len(docling_text):,}")
                    return docling_text
                else:
                    logger.warning(f"⚠️  [EXTRACT] Docling returned empty, falling back to PyPDF")
            else:
                logger.info(f"ℹ️  [EXTRACT] Docling unavailable, using PyPDF")
        else:
            logger.info(f"ℹ️  [EXTRACT] Docling disabled, using PyPDF")

        # 2. PyPDF Fallback mit Fortschrittsanzeige
        try:
            logger.info("=" * 80)
            logger.info(f"📖 [PyPDF] Starting extraction: {filename}")

            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            file_size = os.path.getsize(file_path)

            logger.info(f"   • File size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
            logger.info(f"   • Total pages: {num_pages}")
            logger.info(f"   • Estimated time: ~{num_pages * 0.5:.0f} seconds")
            logger.info("=" * 80)

            text_parts = []
            page_start = time.time()

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Progress logging every 10 pages or last page
                if page_num % 10 == 0 or page_num == num_pages:
                    elapsed = time.time() - page_start
                    avg_time = elapsed / page_num
                    remaining = (num_pages - page_num) * avg_time
                    progress = (page_num / num_pages) * 100

                    logger.info(
                        f"📄 [PyPDF] Progress: {page_num}/{num_pages} pages ({progress:.0f}%) - "
                        f"ETA: {remaining:.0f}s"
                    )

            extracted = "\n".join(text_parts)
            extraction_time = time.time() - extraction_start

            logger.info("=" * 80)
            logger.info(f"✅ [PyPDF] Extraction complete!")
            logger.info(f"   • File: {filename}")
            logger.info(f"   • Duration: {extraction_time:.1f}s")
            logger.info(f"   • Pages: {num_pages}")
            logger.info(f"   • Speed: {extraction_time / num_pages:.2f}s per page")
            logger.info(f"   • Characters: {len(extracted):,}")
            logger.info("=" * 80)

            return extracted

        except Exception as exc:
            extraction_time = time.time() - extraction_start
            logger.error(f"❌ [PyPDF] Extraction failed after {extraction_time:.1f}s: {exc}")
            raise TextExtractionError(f"Failed to extract text from PDF: {exc}")

    @staticmethod
    def extract_metadata(file_path: str) -> Dict[str, Any]:
        """Extrahiert PDF-Metadaten mit Logging"""
        try:
            logger.info(f"📋 [METADATA] Extracting PDF metadata...")
            reader = PdfReader(file_path)
            metadata = reader.metadata

            if metadata:
                result = {
                    "title": metadata.get("/Title", "") or "",
                    "author": metadata.get("/Author", "") or "",
                    "subject": metadata.get("/Subject", "") or "",
                    "creator": metadata.get("/Creator", "") or "",
                    "producer": metadata.get("/Producer", "") or "",
                    "creation_date": str(metadata.get("/CreationDate", "")) or "",
                    "num_pages": len(reader.pages)
                }
                logger.info(f"✅ [METADATA] Extracted: {result.get('num_pages')} pages, "
                            f"title='{result.get('title', 'N/A')}'")
                return result
        except Exception as exc:
            logger.warning(f"⚠️  [METADATA] Extraction failed: {exc}")

        return {"num_pages": 0}

    @staticmethod
    def extract_first_pages(
            file_path: str,
            num_pages: int = 2,
            max_chars: int = 3000
    ) -> str:
        """Extrahiert erste N Seiten für Metadata-Extraktion"""
        logger.info(f"📄 [FIRST PAGES] Extracting first {num_pages} pages (max {max_chars} chars)...")

        # Versuche Docling für bessere Qualität
        docling = DoclingVLMConverter.get_instance()
        if docling:
            logger.info("   → Using Docling for first pages...")
            docling_text = docling.convert(file_path)
            if docling_text:
                result = docling_text[:max_chars]
                logger.info(f"✅ [FIRST PAGES] Extracted {len(result)} chars via Docling")
                return result

        # Fallback zu PyPDF
        try:
            logger.info("   → Using PyPDF for first pages...")
            reader = PdfReader(file_path)
            text_parts = []
            total_chars = 0

            for page in reader.pages[:num_pages]:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    total_chars += len(page_text)
                    if total_chars > max_chars:
                        break

            result = "".join(text_parts)[:max_chars]
            logger.info(f"✅ [FIRST PAGES] Extracted {len(result)} chars via PyPDF")
            return result

        except Exception as exc:
            logger.warning(f"⚠️  [FIRST PAGES] Extraction failed: {exc}")
            return ""


class DOCXExtractor:
    """DOCX Text-Extraktion"""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extrahiert Text aus DOCX mit Logging"""
        filename = os.path.basename(file_path)
        logger.info(f"📄 [DOCX] Extracting from: {filename}")

        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            result = "\n".join(paragraphs)

            logger.info(f"✅ [DOCX] Extracted {len(result):,} characters from {len(paragraphs)} paragraphs")
            return result

        except Exception as exc:
            logger.error(f"❌ [DOCX] Extraction failed: {exc}")
            raise TextExtractionError(f"Failed to extract text from DOCX: {exc}")


class PlainTextExtractor:
    """Plain Text Extraktion"""

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Liest Text-Dateien mit Logging"""
        filename = os.path.basename(file_path)
        logger.info(f"📄 [TEXT] Reading: {filename}")

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                result = f.read()

            logger.info(f"✅ [TEXT] Read {len(result):,} characters")
            return result

        except Exception as exc:
            logger.error(f"❌ [TEXT] Read failed: {exc}")
            raise TextExtractionError(f"Failed to read text file: {exc}")


class FileHandler:
    """Hauptklasse für File-Handling mit verbessertem Logging"""
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extrahiert Text mit automatischer Format-Erkennung"""
        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)

        logger.info(f"📂 [FILE HANDLER] Processing: {filename} (type: {ext})")

        if ext == '.pdf':
            return PDFExtractor.extract_text(file_path)
        elif ext == '.docx':
            return DOCXExtractor.extract_text(file_path)
        elif ext in ['.txt', '.md']:
            return PlainTextExtractor.extract_text(file_path)
        else:
            logger.error(f"❌ [FILE HANDLER] Unsupported file type: {ext}")
            raise UnsupportedFileTypeError(f"Unsupported file type: {ext}")

    @staticmethod
    def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
        """Extrahiert PDF-Metadaten"""
        return PDFExtractor.extract_metadata(file_path)

    @staticmethod
    def extract_first_pages_text(
            file_path: str,
            num_pages: int = 2,
            max_chars: int = 3000
    ) -> str:
        """Extrahiert Text von ersten Seiten"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return PDFExtractor.extract_first_pages(file_path, num_pages, max_chars)
        else:
            full_text = FileHandler.extract_text(file_path)
            return full_text[:max_chars]

    @staticmethod
    def save_upload(file: BinaryIO, filename: str, upload_dir: str) -> str:
        """Speichert hochgeladene Datei"""
        logger.info(f"💾 [UPLOAD] Saving: {filename}")
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, filename)

        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file, buffer)

            file_size = os.path.getsize(file_path)
            logger.info(f"✅ [UPLOAD] Saved: {file_path} ({file_size:,} bytes)")
            return file_path

        except Exception as exc:
            logger.error(f"❌ [UPLOAD] Save failed: {exc}")
            raise FileProcessingError(f"Failed to save uploaded file: {exc}")

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Prüft ob Dateityp unterstützt wird"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in FileHandler.SUPPORTED_EXTENSIONS

    @staticmethod
    def delete_file(file_path: str) -> bool:
        """Löscht Datei vom Filesystem"""
        try:
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"🗑️  [DELETE] Removed: {file_path}")
                return True
            return False
        except Exception as exc:
            logger.warning(f"⚠️  [DELETE] Failed to delete {file_path}: {exc}")
            return False